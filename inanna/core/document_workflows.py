from __future__ import annotations

import csv
import io
import subprocess
import xml.etree.ElementTree as ET
import zipfile
from dataclasses import dataclass, field
from pathlib import Path

from core.desktop_faculty import DesktopFaculty


@dataclass
class DocumentRecord:
    path: str = ""
    title: str = ""
    format: str = ""
    content: str = ""
    word_count: int = 0
    page_count: int = 0
    sheet_names: list[str] = field(default_factory=list)
    error: str | None = None

    @property
    def success(self) -> bool:
        return self.error is None and bool(self.content or self.sheet_names)

    def summary_line(self) -> str:
        parts = [f"doc > {self.format or 'unknown'}"]
        if self.title:
            parts.append(self.title[:50])
        if self.word_count:
            parts.append(f"{self.word_count} words")
        if self.page_count:
            parts.append(f"{self.page_count} pages")
        return " | ".join(parts)


@dataclass
class DocumentWriteResult:
    success: bool
    path: str = ""
    format: str = ""
    word_count: int = 0
    error: str | None = None


@dataclass
class DocumentComprehension:
    title: str = ""
    format: str = ""
    word_count: int = 0
    page_count: int = 0
    summary_lines: list[str] = field(default_factory=list)
    key_points: list[str] = field(default_factory=list)
    suggested_actions: list[str] = field(default_factory=list)

    def to_crown_context(self) -> str:
        lines = [
            f"DOCUMENT: {self.title or '(untitled)'} ({self.format or 'unknown'})",
            f"Size: {self.word_count} words"
            + (f", {self.page_count} pages" if self.page_count > 1 else ""),
        ]
        if self.summary_lines:
            lines.append("SUMMARY:")
            for line in self.summary_lines[:3]:
                lines.append(f"  - {line}")
        if self.key_points:
            lines.append("KEY POINTS:")
            for point in self.key_points[:5]:
                lines.append(f"  - {point}")
        if self.suggested_actions:
            lines.append("SUGGESTED ACTIONS:")
            for action in self.suggested_actions[:3]:
                lines.append(f"  - {action}")
        return "\n".join(lines)


class DocumentDirectReader:
    def read(self, path: str | Path) -> DocumentRecord:
        resolved = Path(path).expanduser().resolve()
        record = DocumentRecord(path=str(resolved), format=resolved.suffix.lower().lstrip("."))
        if not resolved.exists():
            record.error = f"File not found: {resolved}"
            return record

        suffix = resolved.suffix.lower()
        try:
            if suffix in {".txt", ".md", ".rst", ".log"}:
                return self._read_text(resolved, record)
            if suffix == ".docx":
                return self._read_docx(resolved, record)
            if suffix == ".odt":
                return self._read_odt(resolved, record)
            if suffix == ".pdf":
                return self._read_pdf(resolved, record)
            if suffix in {".xlsx", ".xls"}:
                return self._read_xlsx(resolved, record)
            if suffix == ".ods":
                return self._read_ods(resolved, record)
            if suffix == ".csv":
                return self._read_csv(resolved, record)
            record.error = f"Unsupported format: {suffix}"
            return record
        except Exception as exc:
            record.error = str(exc)
            return record

    def _read_text(self, path: Path, record: DocumentRecord) -> DocumentRecord:
        content = path.read_text(encoding="utf-8", errors="replace")
        record.title = path.stem
        record.content = content
        record.word_count = len(content.split())
        record.page_count = max(1, content.count("\f") + 1)
        return record

    def _read_docx(self, path: Path, record: DocumentRecord) -> DocumentRecord:
        import docx

        document = docx.Document(str(path))
        paragraphs = [para.text.strip() for para in document.paragraphs if para.text.strip()]
        record.title = path.stem
        for para in document.paragraphs:
            if para.text.strip() and getattr(para.style, "name", "").startswith("Heading"):
                record.title = para.text.strip()
                break
        record.content = "\n".join(paragraphs)
        record.word_count = len(record.content.split())
        record.page_count = max(1, len(document.sections))
        return record

    def _read_odt(self, path: Path, record: DocumentRecord) -> DocumentRecord:
        try:
            from odf import teletype
            from odf.opendocument import load
            from odf.text import P

            document = load(str(path))
            paragraphs = []
            for node in document.getElementsByType(P):
                text = teletype.extractText(node).strip()
                if text:
                    paragraphs.append(text)
            record.title = path.stem
            record.content = "\n".join(paragraphs)
            record.word_count = len(record.content.split())
            return record
        except Exception:
            with zipfile.ZipFile(path) as archive:
                with archive.open("content.xml") as content_file:
                    tree = ET.parse(content_file)
            texts = [elem.text.strip() for elem in tree.iter() if elem.text and elem.text.strip()]
            record.title = path.stem
            record.content = "\n".join(texts)
            record.word_count = len(record.content.split())
            return record

    def _read_pdf(self, path: Path, record: DocumentRecord) -> DocumentRecord:
        import fitz

        document = fitz.open(str(path))
        record.title = document.metadata.get("title") or path.stem
        record.page_count = len(document)
        pages = [page.get_text() for page in document]
        record.content = "\n\n".join(pages)
        record.word_count = len(record.content.split())
        document.close()
        return record

    def _read_xlsx(self, path: Path, record: DocumentRecord) -> DocumentRecord:
        import openpyxl

        workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        record.title = path.stem
        record.sheet_names = list(workbook.sheetnames)
        lines: list[str] = []
        for sheet_name in workbook.sheetnames:
            worksheet = workbook[sheet_name]
            lines.append(f"=== Sheet: {sheet_name} ===")
            for row in worksheet.iter_rows(max_row=100, values_only=True):
                values = [str(value) if value is not None else "" for value in row]
                row_text = "\t".join(values)
                if row_text.strip():
                    lines.append(row_text)
        workbook.close()
        record.content = "\n".join(lines)
        record.word_count = len(record.content.split())
        return record

    def _read_ods(self, path: Path, record: DocumentRecord) -> DocumentRecord:
        with zipfile.ZipFile(path) as archive:
            with archive.open("content.xml") as content_file:
                tree = ET.parse(content_file)
        table_ns = "{urn:oasis:names:tc:opendocument:xmlns:table:1.0}"
        text_ns = "{urn:oasis:names:tc:opendocument:xmlns:text:1.0}"
        lines: list[str] = []
        for sheet in tree.iter(f"{table_ns}table"):
            name = sheet.get(f"{table_ns}name", "")
            record.sheet_names.append(name)
            lines.append(f"=== Sheet: {name} ===")
            for row in sheet.iter(f"{table_ns}table-row"):
                cells: list[str] = []
                for cell in row.iter(f"{table_ns}table-cell"):
                    text_node = cell.find(f"{text_ns}p")
                    cells.append(text_node.text if text_node is not None and text_node.text else "")
                row_text = "\t".join(cells)
                if row_text.strip():
                    lines.append(row_text)
        record.title = path.stem
        record.content = "\n".join(lines)
        record.word_count = len(record.content.split())
        return record

    def _read_csv(self, path: Path, record: DocumentRecord) -> DocumentRecord:
        content = path.read_text(encoding="utf-8", errors="replace")
        reader = csv.reader(io.StringIO(content))
        rows = list(reader)
        record.title = path.stem
        record.content = "\n".join("\t".join(row) for row in rows[:200])
        record.word_count = sum(len(row) for row in rows)
        return record


class DocumentWriter:
    def write_text(self, path: str | Path, content: str) -> DocumentWriteResult:
        resolved = Path(path).expanduser().resolve()
        try:
            resolved.parent.mkdir(parents=True, exist_ok=True)
            resolved.write_text(content, encoding="utf-8")
            return DocumentWriteResult(
                success=True,
                path=str(resolved),
                format=resolved.suffix.lower().lstrip("."),
                word_count=len(content.split()),
            )
        except Exception as exc:
            return DocumentWriteResult(success=False, path=str(resolved), error=str(exc))

    def write_docx(
        self,
        path: str | Path,
        title: str = "",
        content: str = "",
        paragraphs: list[str] | None = None,
    ) -> DocumentWriteResult:
        import docx

        resolved = Path(path).expanduser().resolve()
        try:
            resolved.parent.mkdir(parents=True, exist_ok=True)
            document = docx.Document()
            if title:
                document.add_heading(title, level=1)
            for paragraph in (paragraphs or content.split("\n\n")):
                paragraph = paragraph.strip()
                if paragraph:
                    document.add_paragraph(paragraph)
            document.save(str(resolved))
            return DocumentWriteResult(
                success=True,
                path=str(resolved),
                format="docx",
                word_count=len((title + "\n" + content).split()),
            )
        except Exception as exc:
            return DocumentWriteResult(success=False, path=str(resolved), error=str(exc))

    def export_pdf_via_libreoffice(
        self,
        input_path: str | Path,
        output_dir: str | Path | None = None,
    ) -> DocumentWriteResult:
        source = Path(input_path).expanduser().resolve()
        if not source.exists():
            return DocumentWriteResult(success=False, path=str(source), error=f"Input file not found: {source}")
        destination = Path(output_dir).expanduser().resolve() if output_dir else source.parent
        command = [
            "soffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(destination),
            str(source),
        ]
        try:
            subprocess.run(command, capture_output=True, text=True, timeout=30)
            pdf_path = destination / f"{source.stem}.pdf"
            if pdf_path.exists():
                return DocumentWriteResult(success=True, path=str(pdf_path), format="pdf")
            return DocumentWriteResult(success=False, path=str(pdf_path), error="PDF not created")
        except subprocess.TimeoutExpired:
            return DocumentWriteResult(success=False, path=str(source), error="LibreOffice timed out during PDF export")
        except FileNotFoundError:
            return DocumentWriteResult(success=False, path=str(source), error="LibreOffice (soffice) not found. Is it installed?")


def build_document_comprehension(record: DocumentRecord) -> DocumentComprehension:
    comprehension = DocumentComprehension(
        title=record.title,
        format=record.format,
        word_count=record.word_count,
        page_count=record.page_count,
    )
    if not record.content:
        return comprehension

    lines = [line.strip() for line in record.content.splitlines() if line.strip()]
    content_lines = [line for line in lines if len(line) > 20 and not line.startswith("#")]
    comprehension.summary_lines = content_lines[:3]
    for line in lines[:200]:
        if len(line) < 120 and (
            line.startswith("#")
            or line.endswith(":")
            or (line[:1].isdigit() and ". " in line[:5])
        ):
            point = line.lstrip("#").strip().rstrip(":")
            if point and len(point) > 5:
                comprehension.key_points.append(point)
        if len(comprehension.key_points) >= 8:
            break
    if record.format == "pdf":
        comprehension.suggested_actions.append("Ask INANNA to summarize specific sections")
    if record.word_count > 2000:
        comprehension.suggested_actions.append("Ask for a shorter summary")
    if record.format in {"xlsx", "ods", "csv"}:
        comprehension.suggested_actions.append("Ask INANNA to explain the data")
    return comprehension


class DocumentWorkflows:
    def __init__(self, desktop: DesktopFaculty) -> None:
        self.desktop = desktop
        self.reader = DocumentDirectReader()
        self.writer = DocumentWriter()

    def read_document(self, path: str) -> tuple[DocumentRecord, DocumentComprehension]:
        record = self.reader.read(path)
        return record, build_document_comprehension(record)

    def write_document(
        self,
        path: str,
        content: str,
        title: str = "",
        format: str = "txt",
    ) -> DocumentWriteResult:
        resolved = Path(path).expanduser()
        suffix = resolved.suffix.lower() or f".{format.lower()}"
        if suffix == ".docx":
            return self.writer.write_docx(resolved, title=title, content=content)
        if title:
            content = f"# {title}\n\n{content}"
        return self.writer.write_text(resolved, content)

    def open_in_libreoffice(self, path: str) -> bool:
        return self.desktop.open_app(f'soffice "{path}"').success

    def export_to_pdf(self, input_path: str, output_dir: str | None = None) -> DocumentWriteResult:
        return self.writer.export_pdf_via_libreoffice(input_path, output_dir)

    def format_read_result(self, record: DocumentRecord, comprehension: DocumentComprehension) -> str:
        if not record.success:
            return f"doc > error: {record.error}"
        return (
            comprehension.to_crown_context()
            + "\n\nCONTENT (first 1500 chars):\n"
            + record.content[:1500]
        )

    def format_write_result(self, result: DocumentWriteResult) -> str:
        if not result.success:
            return f"doc > write error: {result.error}"
        return f"doc > written: {result.path}\nFormat: {result.format} | Words: {result.word_count}"
